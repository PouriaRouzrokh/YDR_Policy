"""
Module for handling document downloads and conversions.
"""
import logging
import os
import re
import urllib.parse
from pathlib import Path
from types import SimpleNamespace

# Document processing libraries
import markdownify
import requests
from docx import Document

# Local imports
from ydrpolicy.data_collection.crawl.processors.llm_processor import \
    process_document_with_ocr
from ydrpolicy.data_collection.crawl.processors.pdf_processor import \
    pdf_to_markdown
from ydrpolicy.data_collection.logger import DataCollectionLogger

# Set up logging
logger = DataCollectionLogger(name="document_processor", level=logging.INFO)

def download_document(url: str, output_dir: str, config: SimpleNamespace) -> str:
    """
    Download a document from a URL and save it to the output directory.
    
    Args:
        url: URL of the document to download
        output_dir: Directory to save the document
        
    Returns:
        Path to the downloaded document
    """
    os.makedirs(output_dir, exist_ok=True)
    
    # Create a filename from the URL
    parsed_url = urllib.parse.urlparse(url)
    
    # For Yale document repository URLs with UUIDs
    if 'files-profile.medicine.yale.edu/documents/' in url:
        # Extract UUID as filename
        match = re.search(r'/documents/([a-f0-9-]+)', parsed_url.path)
        if match:
            filename = f"yale_doc_{match.group(1)}.pdf"  # Assume PDF for Yale documents
        else:
            filename = f"yale_doc_{hash(url) % 10000}.pdf"
    else:
        # Normal filename extraction
        filename = os.path.basename(parsed_url.path)
        
        # If filename is empty or doesn't have an extension, create one
        if not filename or '.' not in filename:
            # Generate a filename based on the URL hash
            filename = f"document_{hash(url) % 10000}{Path(parsed_url.path).suffix}"
            if '.' not in filename:
                # If still no extension, default to .pdf (common for dynamic URLs)
                filename += ".pdf"
    
    file_path = os.path.join(output_dir, filename)
    
    try:
        # Download the file
        logger.info(f"Downloading document from {url}")
        
        # Set up request headers to mimic a browser
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
            'Cache-Control': 'max-age=0'
        }
        
        response = requests.get(url, stream=True, timeout=config.CRAWLER.REQUEST_TIMEOUT, headers=headers)
        response.raise_for_status()
        
        # Check content type to confirm it's a document
        content_type = response.headers.get('Content-Type', '').lower()
        is_document = (
            'pdf' in content_type or 
            'msword' in content_type or 
            'application/vnd.openxmlformats' in content_type or
            'application/vnd.ms-excel' in content_type or
            'application/octet-stream' in content_type
        )
        
        if not is_document:
            logger.warning(f"Content type '{content_type}' may not be a document for {url}")
            # Continue anyway - some servers don't set correct content types
        
        with open(file_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        
        logger.info(f"Document downloaded successfully to {file_path}")
        return file_path
    
    except Exception as e:
        logger.error(f"Error downloading document: {str(e)}")
        return ""

def convert_to_markdown(file_path: str, url: str, config: SimpleNamespace) -> str:
    """
    Convert a document to markdown based on its file type.
    
    Args:
        file_path: Path to the document
        url: Original URL of the document (for OCR fallback)
        
    Returns:
        Document content in markdown format
    """
    file_ext = Path(file_path).suffix.lower()
    
    try:
        # Handle different file types
        if file_ext in ['.pdf']:
            # Always use Mistral OCR for PDFs
            return convert_pdf_to_markdown(file_path, url, config)
        elif file_ext in ['.doc', '.docx']:
            return convert_docx_to_markdown(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_ext}")
            return f"# Unsupported Document\n\nFile: {os.path.basename(file_path)}\nURL: {url}\n\nThis document type is not supported for conversion."
    
    except Exception as e:
        logger.error(f"Error converting document to markdown: {str(e)}")
        return ""

def convert_pdf_to_markdown(file_path: str, url: str, config: SimpleNamespace) -> str:
    """
    Convert a PDF document to markdown using Mistral OCR.
    
    Args:
        file_path: Path to the PDF document
        url: Original URL of the document
        
    Returns:
        PDF content in markdown format
    """
    try:
        logger.info(f"Processing PDF with Mistral OCR: {url}")
        
        # Create a specific output directory for this document
        doc_output_dir = os.path.join(config.PATHS.DOCUMENT_DIR, f"doc_{hash(url) % 10000}")
        os.makedirs(doc_output_dir, exist_ok=True)
        
        # Use the pdf_to_markdown function from pdf_processor
        markdown_path = pdf_to_markdown(url, doc_output_dir)
        
        if markdown_path and os.path.exists(markdown_path):
            with open(markdown_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            # If OCR processing fails, fall back to direct document processing
            logger.warning(f"Mistral OCR processing failed for {url}, trying direct API call")
            markdown_text = process_document_with_ocr(url)
            
            if markdown_text:
                return f"# {os.path.basename(file_path)}\n\nSource: {url}\n\n{markdown_text}"
            else:
                return f"# Failed to Extract Content\n\nFile: {os.path.basename(file_path)}\nURL: {url}\n\nCould not extract content from this PDF."
    
    except Exception as e:
        logger.error(f"Error converting PDF to markdown: {str(e)}")
        return f"# Error Processing PDF\n\nFile: {os.path.basename(file_path)}\nURL: {url}\n\nError: {str(e)}"

def convert_docx_to_markdown(file_path: str) -> str:
    """
    Convert a DOCX document to markdown using python-docx.
    
    Args:
        file_path: Path to the DOCX document
        
    Returns:
        DOCX content in markdown format
    """
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract document title or use filename
        title = os.path.basename(file_path)
        
        # Process each paragraph
        for para in doc.paragraphs:
            if para.text.strip():
                # Handle headings based on style
                style_name = para.style.name.lower()
                if 'heading' in style_name:
                    heading_level = ''.join(filter(str.isdigit, style_name)) or '1'
                    text += f"{'#' * int(heading_level)} {para.text}\n\n"
                else:
                    text += f"{para.text}\n\n"
        
        # Process tables if any
        for table in doc.tables:
            text += "\n| "
            # Add headers
            for cell in table.rows[0].cells:
                text += cell.text + " | "
            text += "\n| "
            
            # Add separator
            for _ in table.rows[0].cells:
                text += "--- | "
            text += "\n"
            
            # Add data rows (skip header)
            for row in table.rows[1:]:
                text += "| "
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
            text += "\n"
        
        return f"# {title}\n\n{text}"
    
    except Exception as e:
        logger.error(f"Error converting DOCX to markdown: {str(e)}")
        return f"# Error Processing DOCX\n\nFile: {os.path.basename(file_path)}\n\nError: {str(e)}"

def html_to_markdown(html_content: str) -> str:
    """
    Convert HTML content to markdown using markdownify.
    
    Args:
        html_content: HTML content to convert
        
    Returns:
        Content in markdown format
    """
    try:
        return markdownify.markdownify(html_content, heading_style="ATX")
    except Exception as e:
        logger.error(f"Error converting HTML to markdown: {str(e)}")
        return ""